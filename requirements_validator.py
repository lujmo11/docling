#!/usr/bin/env python3
"""
Requirements Validation Module

This module validates that all requirements have been properly extracted from
document processing by checking the source document for the highest requirement
number and comparing against extracted requirements.
"""

import logging
import re
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple, Union
import json
import csv

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


logger = logging.getLogger(__name__)


def extract_requirement_numbers_from_docx(docx_path: Union[str, Path]) -> Set[str]:
    """Extract requirement numbers from a DOCX file using python-docx.

    Args:
        docx_path: Path to the DOCX file.

    Returns:
        Set of normalized requirement numbers (e.g., {"1.0", "2.0", "114.0"}).
        Returns empty set if python-docx is unavailable or file cannot be read.
    """
    if not HAS_DOCX:
        logger.warning("python-docx not available, cannot read source DOCX for validation")
        return set()

    requirement_numbers: Set[str] = set()

    try:
        doc = Document(str(docx_path))

        # Paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if text:
                requirement_numbers.update(extract_requirement_numbers_from_text(text))

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    if text:
                        requirement_numbers.update(extract_requirement_numbers_from_text(text))

        logger.info(f"Found {len(requirement_numbers)} requirement numbers in DOCX file")
    except Exception as e:
        logger.warning(f"Failed to read DOCX file directly: {e}")

    return requirement_numbers


def extract_requirement_numbers_from_text(text: str) -> Set[str]:
    """Extract requirement numbers from text using regex patterns.

    Supports:
      #12, #12.0, #12.1, #12.10 etc. We normalize everything to the main integer
      sequence form: e.g. #12.1 -> 12.0

    Args:
        text: Text to search

    Returns:
        Set of requirement numbers normalized to 'N.0'.
    """
    # Match #<number>(.<sub>)? ; capture the main number and optional sub part
    pattern = r'#(\d+)(?:\.(\d+))?\b'
    requirement_numbers: Set[str] = set()
    for match in re.findall(pattern, text):
        main_num = int(match[0])
        requirement_numbers.add(f"{main_num}.0")
    return requirement_numbers


def quick_find_highest_requirement_in_docx(docx_path: Union[str, Path]) -> Optional[int]:
    """Quickly find the highest requirement number by scanning DOCX text backwards.

    This attempts a fast heuristic: extract all text into a list, iterate in
    reverse, and return on the first valid requirement number encountered.
    Falls back to full extraction if heuristic fails.
    """
    if not HAS_DOCX:
        return None
    try:
        doc = Document(str(docx_path))
        lines: list[str] = []
        for p in doc.paragraphs:
            if p.text:
                lines.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        lines.append(cell.text)
        pattern = re.compile(r'#(\d+)(?:\.(\d+))?\b')
        for line in reversed(lines[-500:]):  # limit to last 500 lines for speed
            for m in reversed(list(pattern.finditer(line))):
                try:
                    return int(m.group(1))
                except Exception:
                    continue
        return None
    except Exception:
        return None


def extract_requirement_numbers_from_markdown(md_path: Path) -> Set[str]:
    """
    Extract requirement numbers from a Markdown file.
    
    Args:
        md_path: Path to the Markdown file
        
    Returns:
        Set of requirement numbers found
    """
    requirement_numbers = set()
    
    try:
        text = md_path.read_text(encoding='utf-8')
        requirement_numbers = extract_requirement_numbers_from_text(text)
        logger.debug(f"Found {len(requirement_numbers)} requirement numbers in {md_path.name}")
        return requirement_numbers
        
    except Exception as e:
        logger.warning(f"Failed to read markdown file {md_path}: {e}")
        return set()


def extract_requirement_numbers_from_json(json_path: Path) -> Set[str]:
    """
    Extract requirement numbers from a JSON document file.
    
    Args:
        json_path: Path to the JSON file
        
    Returns:
        Set of requirement numbers found
    """
    requirement_numbers = set()
    
    try:
        with json_path.open('r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Convert to string and search for requirement numbers
        json_text = json.dumps(data, ensure_ascii=False)
        requirement_numbers = extract_requirement_numbers_from_text(json_text)
        logger.debug(f"Found {len(requirement_numbers)} requirement numbers in {json_path.name}")
        return requirement_numbers
        
    except Exception as e:
        logger.warning(f"Failed to read JSON file {json_path}: {e}")
        return set()


def extract_requirement_numbers_from_csv(csv_path: Path) -> Set[str]:
    """
    Extract requirement numbers from a CSV requirements file.
    
    Args:
        csv_path: Path to the CSV file
        
    Returns:
        Set of requirement numbers found
    """
    requirement_numbers: Set[str] = set()
    try:
        with csv_path.open('r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            # Identify requirement id column candidates
            uid_field = None
            for field in reader.fieldnames or []:
                if field.lower() in ("requirement_uid", "requirement_id", "id"):
                    uid_field = field
                    break
            for row in reader:
                # 1) Parse explicit UID column like TPS:057.0 or RS:12.1
                if uid_field and row.get(uid_field):
                    uid_val = row[uid_field].strip()
                    # Match trailing number with optional decimal inside UID (e.g., TPS:057.0)
                    m = re.search(r'(\d+)(?:\.(\d+))?$', uid_val)
                    if m:
                        requirement_numbers.add(f"{int(m.group(1))}.0")
                # 2) Scan key textual fields for inline #numbers (canonical_statement, requirement_raw)
                for field in ("canonical_statement", "requirement_raw"):
                    val = row.get(field)
                    if not val:
                        continue
                    requirement_numbers.update(extract_requirement_numbers_from_text(val))
        logger.debug(f"Found {len(requirement_numbers)} requirement numbers in {csv_path.name} (CSV parsing)")
    except Exception as e:
        logger.warning(f"Failed to read CSV file {csv_path}: {e}")
    return requirement_numbers


def find_highest_requirement_number(requirement_numbers: Set[str]) -> Optional[int]:
    """
    Find the highest requirement number from a set of requirement numbers.
    
    Args:
        requirement_numbers: Set of requirement numbers (e.g., {'1.0', '114.0'})
        
    Returns:
        Highest requirement number as integer, or None if no valid numbers found
    """
    if not requirement_numbers:
        return None
    
    max_num = 0
    for req_num in requirement_numbers:
        try:
            # Extract the integer part (before the decimal)
            num = int(float(req_num))
            max_num = max(max_num, num)
        except (ValueError, TypeError):
            continue
    
    return max_num if max_num > 0 else None


def generate_expected_requirements(max_requirement: int) -> Set[str]:
    """
    Generate the expected set of requirement numbers from 1 to max_requirement.
    
    Args:
        max_requirement: Highest requirement number
        
    Returns:
        Set of expected requirement numbers (e.g., {'1.0', '2.0', ..., '114.0'})
    """
    return {f"{i}.0" for i in range(1, max_requirement + 1)}


def find_missing_requirements(expected: Set[str], found: Set[str]) -> Set[str]:
    """
    Find missing requirements by comparing expected vs found sets.
    
    Args:
        expected: Set of expected requirement numbers
        found: Set of actually found requirement numbers
        
    Returns:
        Set of missing requirement numbers
    """
    return expected - found


def _detect_doc_type(source_docx: Path, output_dir: Path) -> str:
    """Heuristic detection of document type (RS vs TPS)."""
    name = source_docx.name.lower()
    if ' tps ' in f' {name} ' or name.startswith('tps') or ' tps-' in f' {name} ':
        return 'TPS'
    if ' rs ' in f' {name} ' or name.startswith('rs') or ' rs-' in f' {name} ':
        return 'RS'
    # Fallback: inspect requirements CSV doc_meta
    try:
        for csv_path in output_dir.glob('*requirements*.csv'):
            with csv_path.open('r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    meta = row.get('doc_meta')
                    if meta:
                        try:
                            jm = json.loads(meta)
                            dt = (jm.get('document_type') or '').upper()
                            if dt in ('RS', 'TPS'):
                                return dt
                        except Exception:
                            pass
                    break
    except Exception:
        pass
    return 'RS'  # default conservative


def _collect_tps_metrics(source_docx: Path, output_dir: Path) -> Dict:
    """Collect TPS-specific quality metrics without RS sequential missing logic."""
    source_numbers = extract_requirement_numbers_from_docx(source_docx)
    uid_set: Set[str] = set()
    rows = 0
    normative_rows = 0
    norm_re = re.compile(r'\b(shall|must|should|will)\b', re.I)
    prefixed = 0
    hash_pattern = 0
    other_format = 0
    extracted_numbers: Set[str] = set()
    duplicates: Dict[str, int] = {}

    hier_id_re = re.compile(r'^TPS:(\d+(?:\.\d+){1,6})$')
    hier_depths: Dict[int, int] = {}
    hierarchical_ids: Set[str] = set()

    for csv_path in output_dir.glob('*requirements*.csv'):
        try:
            with csv_path.open('r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    rows += 1
                    uid = (row.get('requirement_uid') or '').strip()
                    if uid:
                        if uid in uid_set:
                            duplicates[uid] = duplicates.get(uid, 1) + 1
                        uid_set.add(uid)
                        if uid.upper().startswith('TPS:'):
                            prefixed += 1
                            hm = hier_id_re.match(uid)
                            if hm:
                                hid = hm.group(1)
                                hierarchical_ids.add(hid)
                                depth = hid.count('.') + 1
                                hier_depths[depth] = hier_depths.get(depth, 0) + 1
                        elif re.search(r'#\d+\.\d+$', uid):
                            hash_pattern += 1
                        else:
                            other_format += 1
                        # Extract trailing number pattern for coverage
                        m = re.search(r'(\d+)(?:\.(\d+))?$', uid)
                        if m:
                            extracted_numbers.add(f"{int(m.group(1))}.0")
                    raw = (row.get('requirement_raw') or '') + ' ' + (row.get('canonical_statement') or '')
                    if norm_re.search(raw):
                        normative_rows += 1
        except Exception as e:
            logger.warning(f"Failed reading TPS requirements CSV {csv_path}: {e}")

    coverage_intersection = source_numbers & extracted_numbers if source_numbers else set()
    hier_stats = {
        'total_hier_ids': len(hierarchical_ids),
        'depth_distribution': {str(k): v for k, v in sorted(hier_depths.items())},
        'sample_ids': sorted(list(hierarchical_ids))[:25]
    }
    return {
        'doc_type': 'TPS',
        'source_number_markers': len(source_numbers),
        'extracted_rows': rows,
        'unique_uids': len(uid_set),
        'duplicate_uids': sorted([uid for uid, cnt in duplicates.items() if cnt > 1]),
        'normative_rows': normative_rows,
        'normative_ratio': (normative_rows / rows) if rows else 0.0,
        'uid_format_breakdown': {
            'prefixed_TPS': prefixed,
            'hash_pattern': hash_pattern,
            'other': other_format
        },
        'hierarchical_id_stats': hier_stats,
        'number_marker_coverage': {
            'source_number_markers': len(source_numbers),
            'extracted_number_markers': len(extracted_numbers),
            'covered': len(coverage_intersection),
            'coverage_ratio': (len(coverage_intersection) / len(source_numbers)) if source_numbers else 0.0,
            'uncovered_numbers_sample': sorted(list(source_numbers - coverage_intersection))[:25]
        }
    }


def validate_requirements_extraction(source_docx: Path, output_dir: Path) -> Dict:
    """Validate requirements extraction with separate logic for RS vs TPS.

    RS: original sequential completeness approach.
    TPS: quality & coverage metrics without enforcing linear sequence completeness.
    """
    logger.info(f"Starting requirements validation for {source_docx}")
    doc_type = _detect_doc_type(source_docx, output_dir)

    if doc_type == 'TPS':
        metrics = _collect_tps_metrics(source_docx, output_dir)
        return {
            'status': 'success',
            'source_file': str(source_docx),
            'doc_type': 'TPS',
            'tps_metrics': metrics,
            'processed_files': {}  # kept for interface consistency
        }

    # ---- RS VALIDATION (original path) ----
    source_requirements = extract_requirement_numbers_from_docx(source_docx)
    max_requirement = quick_find_highest_requirement_in_docx(source_docx)
    if max_requirement is None:
        max_requirement = find_highest_requirement_number(source_requirements)
    if max_requirement is None:
        logger.warning("No valid requirement numbers found in source DOCX")
        return {
            'status': 'error',
            'message': 'No valid requirement numbers found in source DOCX',
            'source_file': str(source_docx),
            'source_requirements_found': len(source_requirements),
            'max_requirement': None,
            'doc_type': 'RS'
        }
    logger.info(f"Highest requirement number found: #{max_requirement}.0")
    expected_requirements = generate_expected_requirements(max_requirement)
    results = {
        'status': 'success',
        'source_file': str(source_docx),
        'doc_type': 'RS',
        'max_requirement': max_requirement,
        'expected_count': len(expected_requirements),
        'source_requirements_found': len(source_requirements),
        'source_missing': sorted(find_missing_requirements(expected_requirements, source_requirements)),
        'processed_files': {}
    }
    normalized_docx = source_docx.parent / f"{source_docx.stem}_normalized.docx"
    if normalized_docx.exists():
        norm_requirements = extract_requirement_numbers_from_docx(normalized_docx)
        norm_missing = find_missing_requirements(expected_requirements, norm_requirements)
        results['processed_files']['normalized_docx'] = {
            'file': str(normalized_docx),
            'found_count': len(norm_requirements),
            'missing_count': len(norm_missing),
            'missing_requirements': sorted(norm_missing)
        }
    if output_dir.exists():
        md_files = list(output_dir.glob("*.md"))
        for md_file in md_files:
            md_requirements = extract_requirement_numbers_from_markdown(md_file)
            md_missing = find_missing_requirements(expected_requirements, md_requirements)
            results['processed_files'][f'markdown_{md_file.name}'] = {
                'file': str(md_file),
                'found_count': len(md_requirements),
                'missing_count': len(md_missing),
                'missing_requirements': sorted(md_missing)
            }
        json_files = list(output_dir.glob("*document*.json"))
        for json_file in json_files:
            json_requirements = extract_requirement_numbers_from_json(json_file)
            json_missing = find_missing_requirements(expected_requirements, json_requirements)
            results['processed_files'][f'json_{json_file.name}'] = {
                'file': str(json_file),
                'found_count': len(json_requirements),
                'missing_count': len(json_missing),
                'missing_requirements': sorted(json_missing)
            }
        csv_files = list(output_dir.glob("*requirements*.csv"))
        for csv_file in csv_files:
            csv_requirements = extract_requirement_numbers_from_csv(csv_file)
            csv_missing = find_missing_requirements(expected_requirements, csv_requirements)
            results['processed_files'][f'csv_{csv_file.name}'] = {
                'file': str(csv_file),
                'found_count': len(csv_requirements),
                'missing_count': len(csv_missing),
                'missing_requirements': sorted(csv_missing)
            }
    return results


def save_validation_results(results: Dict, output_dir: Path) -> None:
    """Persist validation results to JSON and text report files."""
    try:
        json_path = output_dir / "requirements_validation.json"
        with json_path.open('w', encoding='utf-8') as f:
            json.dump(results, f, indent=2, ensure_ascii=False)
        # Text report
        report_path = output_dir / "requirements_validation_report.txt"
        from io import StringIO
        buf = StringIO()
        buf.write("Requirements Validation Summary\n")
        if results.get('status') != 'success':
            buf.write(f"Status: {results.get('status')} - {results.get('message','')}\n")
        else:
            doc_type = results.get('doc_type', 'RS')
            buf.write(f"Document Type: {doc_type}\n")
            buf.write(f"Source: {Path(results['source_file']).name}\n")
            if doc_type == 'TPS':
                tm = results.get('tps_metrics', {})
                buf.write(f"Extracted Rows: {tm.get('extracted_rows')}\n")
                buf.write(f"Unique UIDs: {tm.get('unique_uids')}\n")
                buf.write(f"Normative Rows: {tm.get('normative_rows')} ({tm.get('normative_ratio'):.2%})\n")
                cov = tm.get('number_marker_coverage', {})
                buf.write(f"Source Markers: {cov.get('source_number_markers')} Extracted Markers: {cov.get('extracted_number_markers')} Coverage: {cov.get('coverage_ratio'):0.2%}\n")
            else:
                buf.write(f"Highest #: {results['max_requirement']}\n")
                buf.write(f"Expected Count: {results['expected_count']}\n")
                buf.write(f"Source Found: {results['source_requirements_found']}\n")
                src_missing = results.get('source_missing', [])
                buf.write(f"Source Missing Count: {len(src_missing)}\n")
                if src_missing:
                    buf.write(f"Source Missing: {', '.join(src_missing[:25])}{'...' if len(src_missing)>25 else ''}\n")
                for key, data in results.get('processed_files', {}).items():
                    buf.write(f"File: {Path(data['file']).name} Found {data['found_count']} Missing {data['missing_count']}\n")
        report_path.write_text(buf.getvalue(), encoding='utf-8')
    except Exception as e:
        logger.warning(f"Failed to save validation results: {e}")


def print_validation_report(results: Dict) -> None:
    """Print validation report; branches for RS vs TPS."""
    print("\n" + "=" * 80)
    doc_type = results.get('doc_type', 'RS')
    if doc_type == 'TPS':
        print("🧪 TPS REQUIREMENTS QUALITY REPORT")
    else:
        print("📋 REQUIREMENTS VALIDATION REPORT")
    print("=" * 80)

    if results.get('status') != 'success':
        print(f"❌ Error: {results.get('message')}")
        return

    if doc_type == 'TPS':
        source_file = Path(results['source_file']).name
        metrics = results.get('tps_metrics', {})
        print(f"📄 Source File: {source_file}")
        print(f"🔢 Source numbered markers (e.g. #12.5): {metrics.get('source_number_markers')} (used only for coverage)")
        print(f"📦 Extracted rows: {metrics.get('extracted_rows')}")
        print(f"🆔 Unique requirement_uids: {metrics.get('unique_uids')}")
        dup = metrics.get('duplicate_uids') or []
        print(f"♻️ Duplicate UIDs: {len(dup)}{' -> ' + ', '.join(dup[:10]) + ('...' if len(dup) > 10 else '') if dup else ''}")
        print(f"⚖️ Normative rows: {metrics.get('normative_rows')} (ratio {metrics.get('normative_ratio'):.2%})")
        fmt = metrics.get('uid_format_breakdown', {})
        print(f"🔎 UID format breakdown: TPS: {fmt.get('prefixed_TPS')}, #pattern: {fmt.get('hash_pattern')}, other: {fmt.get('other')}")
        hstats = metrics.get('hierarchical_id_stats', {})
        print("\n🪜 Hierarchical ID Stats:")
        print(f"   Total hierarchical IDs: {hstats.get('total_hier_ids')}")
        dd = hstats.get('depth_distribution', {})
        if dd:
            dist_fmt = ', '.join(f"depth {k}: {v}" for k, v in dd.items())
            print(f"   Depth distribution: {dist_fmt}")
        sample_ids = hstats.get('sample_ids', [])
        if sample_ids:
            print(f"   Sample IDs: {', '.join(sample_ids[:15])}{'...' if len(sample_ids)>15 else ''}")
        cov = metrics.get('number_marker_coverage', {})
        print("\n📈 Coverage (document # markers vs extracted):")
        print(f"   Source markers: {cov.get('source_number_markers')} | Extracted markers: {cov.get('extracted_number_markers')}")
        print(f"   Covered markers: {cov.get('covered')} (coverage {cov.get('coverage_ratio'):0.2%})")
        uncovered_sample = cov.get('uncovered_numbers_sample') or []
        if uncovered_sample:
            print(f"   Sample uncovered markers: {', '.join(uncovered_sample)}")
        print("\n✅ Note: TPS report does not list 'missing' sequential requirements because TPS numbering may be non-linear or partial.")
        print("=" * 80)
        return

    # RS branch (original formatting)
    source_file = Path(results['source_file']).name
    max_req = results['max_requirement']
    expected_count = results['expected_count']
    print(f"📄 Source File: {source_file}")
    print(f"🎯 Highest Requirement: #{max_req}.0")
    print(f"📊 Expected Requirements: #1.0 to #{max_req}.0 ({expected_count} total)")
    source_found = results['source_requirements_found']
    source_missing = results['source_missing']
    source_missing_count = len(source_missing)
    print(f"\n🔍 SOURCE FILE ANALYSIS:")
    print(f"   ✅ Found: {source_found}/{expected_count} requirements")
    if source_missing_count > 0:
        print(f"   ❌ Missing: {source_missing_count} requirements")
        if source_missing_count <= 10:
            print(f"   📝 Missing numbers: {', '.join(f'#{req}' for req in source_missing)}")
        else:
            print(f"   📝 Missing numbers: {', '.join(f'#{req}' for req in source_missing[:10])}... (+{source_missing_count-10} more)")
    else:
        print(f"   ✅ All requirements found in source!")
    print(f"\n📁 PROCESSED FILES ANALYSIS:")
    if not results['processed_files']:
        print("   ⚠️  No processed files found to validate")
        return
    for file_key, file_results in results['processed_files'].items():
        file_name = Path(file_results['file']).name
        found_count = file_results['found_count']
        missing_count = file_results['missing_count']
        missing_reqs = file_results['missing_requirements']
        file_type = file_key.split('_')[0].upper()
        status = '✅' if missing_count == 0 else '⚠️'
        print(f"\n   {status} {file_type}: {file_name}")
        print(f"      Found: {found_count}/{expected_count} requirements")
        if missing_count > 0:
            if missing_count <= 10:
                print(f"      Missing numbers: {', '.join(f'#{req}' for req in missing_reqs)}")
            else:
                print(f"      Missing numbers: {', '.join(f'#{req}' for req in missing_reqs[:10])}... (+{missing_count-10} more)")
        else:
            print(f"      ✅ All requirements captured!")
    print(f"\n📈 SUMMARY:")
    total_files = len(results['processed_files']) + 1
    perfect_files = sum(1 for f in results['processed_files'].values() if f['missing_count'] == 0)
    if source_missing_count == 0:
        perfect_files += 1
    print(f"   📄 Files analyzed: {total_files}")
    print(f"   ✅ Perfect extractions: {perfect_files}/{total_files}")
    print(f"   🎯 Overall success rate: {perfect_files/total_files*100:.1f}%")
    print("=" * 80)


if __name__ == "__main__":
    # Simple CLI for testing
    import argparse
    import sys
    
    def setup_logging(level: int = logging.INFO) -> None:
        logging.basicConfig(
            level=level,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            handlers=[logging.StreamHandler(sys.stdout)]
        )
    
    parser = argparse.ArgumentParser(description="Validate requirements extraction")
    parser.add_argument("source_docx", help="Path to source DOCX file")
    parser.add_argument("output_dir", help="Path to output directory")
    parser.add_argument("--debug", action="store_true", help="Enable debug logging")
    
    args = parser.parse_args()
    
    setup_logging(logging.DEBUG if args.debug else logging.INFO)
    
    try:
        results = validate_requirements_extraction(Path(args.source_docx), Path(args.output_dir))
        print_validation_report(results)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)