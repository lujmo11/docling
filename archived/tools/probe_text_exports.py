import re
import json
from pathlib import Path

DOCX_PATH = Path("A012-5599 TPS 4.8 MW Air Cooler IG GEN V05 - CAC.docx")
ID_PATTERN = re.compile(r"\b\d+(?:\.\d+){1,6}\b")

def scan(text: str):
    ids = sorted(set(ID_PATTERN.findall(text)))
    depth4 = [i for i in ids if i.count('.') >= 3]
    target_block = [i for i in ids if i.startswith('4.1.1.')]  # specific family we need
    return {
        "count": len(ids),
        "sample": ids[:12],
        "depth>=4_count": len(depth4),
        "sample_depth>=4": depth4[:8],
        "4.1.1.x_present": bool(target_block),
        "4.1.1.x_samples": target_block[:8]
    }

def try_python_docx():
    try:
        import docx  # type: ignore
    except ImportError:
        return {"ok": False, "error": "python-docx not installed"}
    try:
        d = docx.Document(str(DOCX_PATH))
        parts = []
        for p in d.paragraphs:
            parts.append(p.text)
        # include table cell paragraphs explicitly
        for t in d.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        parts.append(p.text)
        text = "\n".join(parts)
        return {"ok": True, "engine": "python-docx", **scan(text)}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def try_aspose():
    try:
        import aspose.words as aw  # type: ignore
    except ImportError:
        return {"ok": False, "error": "aspose.words not installed"}
    try:
        doc = aw.Document(str(DOCX_PATH))
        txt = doc.get_text()  # full document text layout-like
        return {"ok": True, "engine": "aspose", **scan(txt)}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def try_docx2python():
    try:
        from docx2python import docx2python  # type: ignore
    except ImportError:
        return {"ok": False, "error": "docx2python not installed"}
    try:
        result = docx2python(str(DOCX_PATH))
        # Flatten body
        def flatten(node):
            if isinstance(node, (list, tuple)):
                for x in node:
                    yield from flatten(x)
            else:
                if isinstance(node, str):
                    yield node
        text = "\n".join(x for x in flatten(result.body) if x)
        return {"ok": True, "engine": "docx2python", **scan(text)}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def try_docling():
    try:
        from docling.document_converter import DocumentConverter
    except ImportError:
        return {"ok": False, "error": "docling not installed"}
    try:
        conv = DocumentConverter()
        res = conv.convert(str(DOCX_PATH))
        doc = getattr(res, 'document', None)
        if not doc:
            return {"ok": False, "error": "no doc"}
        if hasattr(doc, 'export_to_markdown'):
            md = doc.export_to_markdown()
        else:
            md = str(doc)
        return {"ok": True, "engine": "docling-markdown", **scan(md)}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def main():
    if not DOCX_PATH.exists():
        print(json.dumps({"error": f"Missing file {DOCX_PATH}"}, indent=2))
        return
    results = []
    # Order: aspose, python-docx, docx2python, docling
    results.append(try_aspose())
    results.append(try_python_docx())
    results.append(try_docx2python())
    results.append(try_docling())
    mapped = {}
    for r in results:
        key = r.get('engine','?')
        mapped[key] = r
    print(json.dumps(mapped, indent=2))

if __name__ == "__main__":
    main()